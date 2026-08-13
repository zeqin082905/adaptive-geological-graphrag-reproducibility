"""
步骤1-2：地质资料解析与数据清洗
支持 PDF（含扫描件 OCR）、Word(.docx)、纯文本，统一输出按页组织的结构化文本。
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ── 数据结构 ──────────────────────────────────────────────────────────────────
@dataclass
class PageText:
    """单页文本的标准化表示。"""
    doc_id: str                 # 文档唯一标识（文件名 stem）
    page_num: int               # 页码（从 1 起）
    raw_text: str               # 原始文本
    cleaned_text: str = ""      # 清洗后文本
    source_path: str = ""       # 原始文件路径
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """解析后的完整文档。"""
    doc_id: str
    source_path: str
    total_pages: int
    pages: list[PageText] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n".join(p.cleaned_text for p in self.pages)


# ── 解析器基类 ────────────────────────────────────────────────────────────────
class BaseParser:
    def parse(self, file_path: Path) -> ParsedDocument:
        raise NotImplementedError


# ── PDF 解析器 ────────────────────────────────────────────────────────────────
class PDFParser(BaseParser):
    """
    优先使用 PyMuPDF（fitz）提取文本；
    若页面文字量极少（扫描件），自动降级到 PaddleOCR。
    """

    OCR_FALLBACK_THRESHOLD = 50  # 字符数低于此值则触发 OCR

    def parse(self, file_path: Path) -> ParsedDocument:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("请安装 PyMuPDF: pip install pymupdf")

        doc_id = file_path.stem
        pdf = fitz.open(str(file_path))
        pages: list[PageText] = []

        for page_idx in range(len(pdf)):
            page = pdf[page_idx]
            text = page.get_text("text")

            # 扫描件降级 OCR
            if len(text.strip()) < self.OCR_FALLBACK_THRESHOLD:
                logger.info(f"[OCR] {doc_id} 第{page_idx+1}页文字稀少，尝试 OCR")
                text = self._ocr_page(page)

            pages.append(PageText(
                doc_id=doc_id,
                page_num=page_idx + 1,
                raw_text=text,
                source_path=str(file_path),
            ))

        pdf.close()
        return ParsedDocument(
            doc_id=doc_id,
            source_path=str(file_path),
            total_pages=len(pages),
            pages=pages,
        )

    def _ocr_page(self, page) -> str:
        """使用 PaddleOCR 对单页图像进行文字识别。"""
        try:
            from paddleocr import PaddleOCR
            import numpy as np
            ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
            pix = page.get_pixmap(dpi=200)
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                pix.height, pix.width, pix.n
            )
            result = ocr.ocr(img, cls=True)
            lines = [line[1][0] for block in result for line in block] if result else []
            return "\n".join(lines)
        except ImportError:
            logger.warning("PaddleOCR 未安装，跳过 OCR，返回空文本")
            return ""


# ── Word 解析器 ───────────────────────────────────────────────────────────────
class DocxParser(BaseParser):
    def parse(self, file_path: Path) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc_id = file_path.stem
        docx = DocxDocument(str(file_path))
        full_text = "\n".join(p.text for p in docx.paragraphs if p.text.strip())

        # Word 文档视为单页
        page = PageText(
            doc_id=doc_id,
            page_num=1,
            raw_text=full_text,
            source_path=str(file_path),
        )
        return ParsedDocument(
            doc_id=doc_id,
            source_path=str(file_path),
            total_pages=1,
            pages=[page],
        )


# ── 数据清洗器 ────────────────────────────────────────────────────────────────
class GeoCleaner:
    """
    步骤2：针对地质报告的专用清洗流水线。
    去噪 → 去页眉页脚 → 去目录/图表干扰。
    """

    # 页眉页脚特征模式（可根据实际语料扩展）
    HEADER_FOOTER_PATTERNS = [
        r"^\s*[-—]{3,}\s*$",                          # 纯分隔线
        r"^\s*第\s*\d+\s*页\s*(?:共\s*\d+\s*页)?\s*$",  # 页码行
        r"^\s*\d+\s*/\s*\d+\s*$",                      # n/N 页码
        r"^\s*[A-Z]{2,}-\d{4}-\d+\s*$",               # 报告编号行
    ]

    # 目录行特征（中文目录）
    TOC_LINE_PATTERN = re.compile(
        r"^.{1,30}[\u4e00-\u9fff\w\s]{0,20}\.{3,}\s*\d+\s*$"
    )

    def clean(self, doc: ParsedDocument) -> ParsedDocument:
        """对文档所有页面执行清洗，原地修改 cleaned_text 字段。"""
        header_footer_re = [re.compile(p, re.MULTILINE) for p in self.HEADER_FOOTER_PATTERNS]

        for page in doc.pages:
            text = page.raw_text
            text = self._remove_header_footer(text, header_footer_re)
            text = self._remove_toc_lines(text)
            text = self._remove_noise(text)
            page.cleaned_text = text.strip()

        return doc

    def _remove_header_footer(self, text: str, patterns: list) -> str:
        lines = text.split("\n")
        cleaned = []
        for line in lines:
            if any(p.match(line.strip()) for p in patterns):
                continue
            cleaned.append(line)
        return "\n".join(cleaned)

    def _remove_toc_lines(self, text: str) -> str:
        lines = text.split("\n")
        return "\n".join(
            line for line in lines if not self.TOC_LINE_PATTERN.match(line.strip())
        )

    def _remove_noise(self, text: str) -> str:
        # 合并多余空行
        text = re.sub(r"\n{3,}", "\n\n", text)
        # 去除控制字符（保留换行、制表符）
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        return text


# ── 文档库加载入口 ────────────────────────────────────────────────────────────
class CorpusLoader:
    """
    步骤1：扫描指定目录，解析所有支持格式的地质资料文档。
    """
    SUPPORTED_SUFFIXES = {".pdf": PDFParser, ".docx": DocxParser, ".txt": None}

    def __init__(self):
        self.cleaner = GeoCleaner()

    def load_directory(self, corpus_dir: Path) -> list[ParsedDocument]:
        """加载整个文档库目录，返回清洗后的文档列表。"""
        corpus_dir = Path(corpus_dir)
        docs: list[ParsedDocument] = []

        for file_path in sorted(corpus_dir.rglob("*")):
            suffix = file_path.suffix.lower()
            if suffix not in self.SUPPORTED_SUFFIXES:
                continue

            logger.info(f"解析文档: {file_path.name}")
            try:
                doc = self._parse_file(file_path, suffix)
                doc = self.cleaner.clean(doc)
                docs.append(doc)
            except Exception as e:
                logger.error(f"解析失败 [{file_path.name}]: {e}")

        logger.info(f"共加载 {len(docs)} 份文档，来源目录: {corpus_dir}")
        return docs

    def _parse_file(self, file_path: Path, suffix: str) -> ParsedDocument:
        if suffix == ".pdf":
            return PDFParser().parse(file_path)
        elif suffix == ".docx":
            return DocxParser().parse(file_path)
        elif suffix == ".txt":
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            page = PageText(
                doc_id=file_path.stem,
                page_num=1,
                raw_text=text,
                source_path=str(file_path),
            )
            doc = ParsedDocument(
                doc_id=file_path.stem,
                source_path=str(file_path),
                total_pages=1,
                pages=[page],
            )
            doc.pages[0].cleaned_text = text
            return doc
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")
